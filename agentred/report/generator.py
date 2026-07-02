"""Report generator: produces markdown and DOCX reports from run results."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from agentred.schemas import RunReport


def generate_markdown(report: RunReport) -> str:
    """Generate a markdown report from a run.

    Args:
        report: The RunReport to render.

    Returns:
        Markdown string.
    """
    r = report.reduce_result
    lines: list[str] = []

    lines.append(f"# AgentRed Demand Intelligence Report")
    lines.append(f"")

    # Summary
    lines.append(f"## Summary")
    lines.append(f"")
    lines.append(f"| Metric | Value |")
    lines.append(f"|---|---|")
    lines.append(f"| Run ID | `{report.run_id}` |")
    lines.append(f"| Sources | {', '.join(report.sources_used)} |")
    lines.append(f"| Posts collected | {report.total_posts_collected} |")
    lines.append(f"| Posts analyzed | {report.total_posts_after_dedup} |")
    lines.append(f"| Signals extracted | {r.total_signals_extracted if r else 0} |")
    lines.append(f"| Opportunities ranked | {len(r.opportunities) if r else 0} |")
    lines.append(f"| Wall clock | {report.wall_clock_seconds:.1f}s |")
    lines.append(f"")

    if not r or not r.opportunities:
        lines.append("_No opportunities were identified in this run._")
        return "\n".join(lines)

    # Top opportunities
    lines.append(f"## Top Build Opportunities")
    lines.append(f"")

    for opp in r.opportunities:
        lines.append(f"### {opp.rank}. {opp.title}")
        lines.append(f"")
        lines.append(f"**What to build:** {opp.what_to_build}")
        lines.append(f"")
        lines.append(f"| Attribute | Value |")
        lines.append(f"|---|---|")
        lines.append(f"| Demand score | {opp.demand_score:.1f}/100 |")
        lines.append(f"| Automation potential | {opp.automation_potential:.0%} |")
        lines.append(f"| Competition | {opp.competition_level} |")
        lines.append(f"| Budget range | {opp.estimated_budget_range or 'N/A'} |")
        lines.append(f"| Source count | {opp.source_count} |")
        lines.append(f"| Skills | {', '.join(opp.skill_signals) or 'N/A'} |")
        lines.append(f"")
        if opp.why_now:
            lines.append(f"**Why now:** {opp.why_now}")
            lines.append(f"")
        if opp.evidence:
            lines.append(f"**Evidence:**")
            lines.append(f"")
            for url in opp.evidence[:5]:
                lines.append(f"- {url}")
            lines.append(f"")

    # Skill frequency
    lines.append(f"## Skill Frequency")
    lines.append(f"")
    lines.append(f"| Skill | Mentions |")
    lines.append(f"|---|---|")
    for skill, count in sorted(
        r.skill_frequency.items(), key=lambda x: x[1], reverse=True
    )[:20]:
        lines.append(f"| {skill} | {count} |")
    lines.append(f"")

    # Category distribution
    if r.category_distribution:
        lines.append(f"## Category Distribution")
        lines.append(f"")
        lines.append(f"| Category | Count |")
        lines.append(f"|---|---|")
        for cat, count in sorted(
            r.category_distribution.items(), key=lambda x: x[1], reverse=True
        ):
            lines.append(f"| {cat} | {count} |")
        lines.append(f"")

    # Contradictions
    if r.contradictions:
        lines.append(f"## Contradictions & Gaps")
        lines.append(f"")
        for c in r.contradictions:
            lines.append(f"- {c}")
        lines.append(f"")

    # Confidence notes
    if r.confidence_notes:
        lines.append(f"## Confidence Notes")
        lines.append(f"")
        lines.append(f"{r.confidence_notes}")
        lines.append(f"")

    return "\n".join(lines)


def generate_docx(report: RunReport, output_path: str | Path) -> Path:
    """Generate a DOCX report from a run.

    Requires python-docx: pip install agentred[report]

    Args:
        report: The RunReport to render.
        output_path: Where to save the .docx file.

    Returns:
        Path to the saved file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX output. "
            "Install with: pip install agentred[report]"
        )

    doc = Document()
    r = report.reduce_result

    # Title
    doc.add_heading("AgentRed Demand Intelligence Report", level=0)

    # Summary table
    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Light Shading"
    cells = summary_table.rows[0].cells
    cells[0].text = "Metric"
    cells[1].text = "Value"

    rows = [
        ("Run ID", report.run_id),
        ("Sources", ", ".join(report.sources_used)),
        ("Posts collected", str(report.total_posts_collected)),
        ("Posts analyzed", str(report.total_posts_after_dedup)),
        ("Signals extracted", str(r.total_signals_extracted if r else 0)),
        ("Opportunities", str(len(r.opportunities) if r else 0)),
        ("Wall clock", f"{report.wall_clock_seconds:.1f}s"),
    ]
    for label, value in rows:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = value

    if not r or not r.opportunities:
        doc.add_paragraph("No opportunities were identified in this run.")
        return _save_docx(doc, output_path)

    # Opportunities
    doc.add_heading("Top Build Opportunities", level=1)

    for opp in r.opportunities:
        doc.add_heading(f"{opp.rank}. {opp.title}", level=2)
        doc.add_paragraph(opp.what_to_build)

        # Detail table
        opp_table = doc.add_table(rows=1, cols=2)
        opp_table.style = "Light Shading"
        opp_table.rows[0].cells[0].text = "Attribute"
        opp_table.rows[0].cells[1].text = "Value"

        detail_rows = [
            ("Demand score", f"{opp.demand_score:.1f}/100"),
            ("Automation potential", f"{opp.automation_potential:.0%}"),
            ("Competition", opp.competition_level),
            ("Budget range", opp.estimated_budget_range or "N/A"),
            ("Source count", str(opp.source_count)),
            ("Skills", ", ".join(opp.skill_signals) or "N/A"),
        ]
        for label, value in detail_rows:
            row = opp_table.add_row().cells
            row[0].text = label
            row[1].text = value

        if opp.why_now:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Why now: ")
            run.bold = True
            p.add_run(opp.why_now)

        if opp.evidence:
            doc.add_paragraph("Evidence:")
            for url in opp.evidence[:5]:
                doc.add_paragraph(url, style="List Bullet")

        doc.add_paragraph()  # spacer

    # Skill frequency
    doc.add_heading("Skill Frequency", level=1)
    skill_table = doc.add_table(rows=1, cols=2)
    skill_table.style = "Light Shading"
    skill_table.rows[0].cells[0].text = "Skill"
    skill_table.rows[0].cells[1].text = "Mentions"

    for skill, count in sorted(
        r.skill_frequency.items(), key=lambda x: x[1], reverse=True
    )[:20]:
        row = skill_table.add_row().cells
        row[0].text = skill
        row[1].text = str(count)

    # Confidence notes
    if r.confidence_notes:
        doc.add_heading("Confidence Notes", level=1)
        doc.add_paragraph(r.confidence_notes)

    return _save_docx(doc, output_path)


def _save_docx(doc, path: str | Path) -> Path:
    """Save a python-docx Document to the given path."""
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return p
